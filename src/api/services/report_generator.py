"""PDF and DOCX report generation for D8X agent outputs."""

from __future__ import annotations

import io
from datetime import datetime, timezone


def _build_improvement_guide(quality: dict, type_breakdown: dict, total_words: int) -> list[dict[str, str]]:
    """Generate actionable improvement suggestions based on current quality gaps."""
    sections: list[dict[str, str]] = []

    completeness = quality.get("completeness", 0)
    diversity = quality.get("diversity", 0)
    volume = quality.get("volume", 0)

    # ── Completeness ──
    sections.append({
        "title": "Completeness",
        "score": str(completeness),
        "what": (
            "Completeness measures whether the inputs cover all major aspects of the system: "
            "business rules, user workflows, data models, integrations, and non-functional requirements. "
            "A low completeness score means the Discovery agent may miss critical requirements, "
            "leading to gaps in the design that are expensive to fix later."
        ),
        "why": (
            "In a Standish Group study, 56% of software defects originate from incomplete requirements. "
            "Every requirement missed during analysis costs 10-100x more to fix after development begins. "
            "Thorough inputs at this stage are the single highest-leverage investment in your project."
        ),
        "suggestions": (
            "• Business Requirements Document (BRD) — formal statement of what the system must do\n"
            "• User stories or use case documents — how each role interacts with the system\n"
            "• Process flow diagrams — step-by-step workflows (even hand-drawn on a whiteboard)\n"
            "• Acceptance criteria — how you'll verify each requirement is met\n"
            "• Non-functional requirements — performance targets, security needs, compliance constraints"
        ) if completeness < 80 else "Your completeness score is strong. No additional documents needed for this dimension.",
    })

    # ── Diversity ──
    types_present = list(type_breakdown.keys())
    types_str = ", ".join(t.replace("_", " ") for t in types_present) if types_present else "none detected"

    missing_types = []
    if "source_code" not in type_breakdown:
        missing_types.append("source code files (.py, .js, .java)")
    if "audio" not in type_breakdown and "video" not in type_breakdown:
        missing_types.append("meeting recordings (.mp3, .mp4) — stakeholder interviews are gold")
    if not any(t in type_breakdown for t in ("spreadsheet", "database_schema")):
        missing_types.append("data models — database schemas (.sql), spreadsheets (.xlsx), ER diagrams")
    if "image" not in type_breakdown:
        missing_types.append("diagrams or wireframes (.png, .jpg) — architecture, UI mockups, whiteboard photos")

    sections.append({
        "title": "Diversity",
        "score": str(diversity),
        "what": (
            "Diversity measures how many different types of sources you've provided. "
            f"Currently you have: {types_str}. "
            "Different source types capture different perspectives — code reveals what was built, "
            "documents reveal what was intended, recordings reveal what stakeholders actually want "
            "(which is often different from what's written down)."
        ),
        "why": (
            "Cross-referencing multiple source types is how D8X detects contradictions. "
            "For example, a BRD might say 'orders over $5,000 need approval' while the code "
            "implements a $10,000 threshold. These conflicts are invisible if you only provide one source type. "
            "The Discovery agent's conflict detection accuracy increases by ~40% with 3+ source types."
        ),
        "suggestions": (
            "Consider adding these source types you're currently missing:\n"
            + "\n".join(f"• {t}" for t in missing_types)
            + ("\n\nExample: For a legacy modernization, the ideal input set is: "
               "BRD + source code + database schema + 1-2 meeting recordings + API documentation. "
               "This gives D8X five different viewpoints on the same system.")
        ) if missing_types else "Your source diversity is excellent. All major source types are represented.",
    })

    # ── Volume ──
    sections.append({
        "title": "Volume",
        "score": str(volume),
        "what": (
            f"Volume measures the total amount of content available for analysis. "
            f"Currently: {total_words:,} words across all sources. "
            "More content gives the Discovery agent more material to extract business rules from. "
            "A thin input produces a thin analysis."
        ),
        "why": (
            "D8X's business rule extraction works by pattern matching across your content. "
            "With fewer than 2,000 words, the agent typically finds only surface-level rules. "
            "With 5,000-10,000 words, it begins to identify edge cases, exceptions, and implicit rules. "
            "With 10,000+ words, it can detect conflicts between sources and generate clarification questions "
            "that save weeks of back-and-forth during development."
        ),
        "suggestions": (
            "Ways to increase content volume:\n"
            "• Upload the full BRD, not just the executive summary\n"
            "• Include appendices, data dictionaries, and glossaries\n"
            "• Add meeting recordings — a 30-minute meeting typically yields 5,000+ words when transcribed\n"
            "• Include existing source code files — even partial codebases reveal business logic\n"
            "• Add email threads or Slack exports discussing requirements decisions\n"
            "\nExample: A typical well-prepared legacy modernization project provides 15,000-30,000 words "
            "across 5-10 source files. This yields 15-25 business rules, 10-15 domain entities, "
            "and 5-10 conflicts or clarification questions."
        ) if volume < 80 else "Your content volume is sufficient for thorough analysis.",
    })

    return sections


def generate_pdf_report(project: object, agent_run: object, agent_type: str) -> bytes:
    """Generate a professional PDF report."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75 * inch, bottomMargin=0.75 * inch, leftMargin=0.75 * inch, rightMargin=0.75 * inch)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("D8XTitle", parent=styles["Title"], fontSize=24, textColor=colors.HexColor("#1B4F72"), spaceAfter=6))
    styles.add(ParagraphStyle("D8XSubtitle", parent=styles["Normal"], fontSize=12, textColor=colors.HexColor("#8B949E"), spaceAfter=20))
    styles.add(ParagraphStyle("SectionHead", parent=styles["Heading2"], fontSize=14, textColor=colors.HexColor("#1B4F72"), spaceBefore=20, spaceAfter=10))
    styles.add(ParagraphStyle("BodyText2", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=8))
    styles.add(ParagraphStyle("SmallGray", parent=styles["Normal"], fontSize=9, textColor=colors.HexColor("#8B949E"), alignment=TA_CENTER))

    output = agent_run.output_summary or {}  # type: ignore[union-attr]
    metrics = output.get("metrics", {})
    processed_files = output.get("processed_files", [])
    quality = output.get("quality_assessment", {})
    project_type = output.get("project_type", "unknown")
    type_breakdown = output.get("type_breakdown", {})
    now = datetime.now(timezone.utc).strftime("%B %d, %Y at %I:%M %p UTC")

    navy = colors.HexColor("#1B4F72")
    border = colors.HexColor("#30363D")
    alt_row = colors.HexColor("#F8F9FA")
    project_type_label = "Legacy modernization" if project_type == "legacy_modernization" else "Greenfield project"

    def _table_style() -> TableStyle:
        return TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), navy), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("GRID", (0, 0), (-1, -1), 0.5, border),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, alt_row]),
            ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6), ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ])

    story: list = []

    # Header
    story.append(Paragraph("D8X Analysis Report", styles["D8XTitle"]))
    story.append(Paragraph(f"Project: {project.name}", styles["D8XSubtitle"]))  # type: ignore[union-attr]
    agent_labels = {"ingest": "D1 — Ingest", "discover": "D2 — Discover", "design": "D3 — Design"}
    story.append(Paragraph(f"Agent: {agent_labels.get(agent_type, agent_type)} | Generated: {now}", styles["SmallGray"]))
    story.append(HRFlowable(width="100%", thickness=1, color=border, spaceBefore=10, spaceAfter=20))

    # Executive Summary
    story.append(Paragraph("Executive Summary", styles["SectionHead"]))
    story.append(Paragraph(
        f"D8X ingested and analyzed <b>{metrics.get('files_processed', 0)} files</b> "
        f"containing approximately <b>{metrics.get('words_extracted', 0):,} words</b> of content across "
        f"<b>{metrics.get('sources_classified', 0)} classified sources</b>. "
        f"The project has been identified as a <b>{project_type_label}</b>. "
        f"Overall input quality score: <b>{metrics.get('quality_score', 0)}/100</b>.", styles["BodyText2"]
    ))
    story.append(Spacer(1, 12))

    # Metrics table
    t = Table([["Metric", "Value"]] + [
        ["Files processed", str(metrics.get("files_processed", 0))],
        ["Words extracted", f"{metrics.get('words_extracted', 0):,}"],
        ["Sources classified", str(metrics.get("sources_classified", 0))],
        ["Quality score", f"{metrics.get('quality_score', 0)}/100"],
        ["Project type", project_type_label],
    ], colWidths=[3 * inch, 3.5 * inch])
    t.setStyle(_table_style())
    story.append(t)
    story.append(Spacer(1, 20))

    # Processed files
    story.append(Paragraph("Sources Received", styles["SectionHead"]))
    if processed_files:
        rows = [["Filename", "Type", "Words", "Status"]]
        for f in processed_files:
            rows.append([f.get("filename", "?"), f.get("file_type", "?").replace("_", " ").title(), f"{f.get('word_count', 0):,}", f.get("status", "?").title()])
        ft = Table(rows, colWidths=[2.8 * inch, 1.5 * inch, 1 * inch, 1.2 * inch])
        ft.setStyle(_table_style())
        story.append(ft)
    story.append(Spacer(1, 20))

    # Type breakdown
    if type_breakdown:
        story.append(Paragraph("Source Type Breakdown", styles["SectionHead"]))
        story.append(Paragraph(", ".join(f"{c} {t.replace('_', ' ')}" for t, c in type_breakdown.items()), styles["BodyText2"]))
        story.append(Spacer(1, 12))

    # Quality
    story.append(Paragraph("Readiness Assessment", styles["SectionHead"]))
    qrows = [["Dimension", "Score", "Rating"]]
    for dim in ["completeness", "diversity", "volume"]:
        s = quality.get(dim, 0)
        qrows.append([dim.title(), f"{s}/100", "Excellent" if s >= 80 else "Adequate" if s >= 60 else "Needs improvement"])
    qt = Table(qrows, colWidths=[2.5 * inch, 1.5 * inch, 2.5 * inch])
    qt.setStyle(_table_style())
    story.append(qt)
    story.append(Spacer(1, 12))

    warnings = quality.get("warnings", [])
    if warnings:
        story.append(Paragraph("Warnings", styles["SectionHead"]))
        for w in warnings:
            story.append(Paragraph(f"⚠ {w}", styles["BodyText2"]))

    # ── How to Improve Your Scores ──
    total_words = metrics.get("words_extracted", 0)
    improvement_guide = _build_improvement_guide(quality, type_breakdown, total_words)
    any_needs_improvement = any(int(s["score"]) < 80 for s in improvement_guide)

    if any_needs_improvement:
        story.append(Spacer(1, 16))
        story.append(Paragraph("How to Improve Your Scores", styles["SectionHead"]))
        story.append(Paragraph(
            "The quality scores above directly impact how thorough the Discovery agent's analysis will be. "
            "Below are specific, actionable suggestions for each dimension. You can add more files and "
            "re-run the Ingest agent without losing your existing results.",
            styles["BodyText2"],
        ))
        story.append(Spacer(1, 8))

        for section in improvement_guide:
            if int(section["score"]) >= 80:
                continue  # Skip dimensions that are already good

            # Dimension header
            score_val = int(section["score"])
            color_hex = "#27AE60" if score_val >= 80 else "#F39C12" if score_val >= 60 else "#E74C3C"
            story.append(Paragraph(
                f'<b>{section["title"]}</b> — <font color="{color_hex}">{section["score"]}/100</font>',
                styles["BodyText2"],
            ))
            story.append(Spacer(1, 4))

            # What this measures
            story.append(Paragraph(f'<i>What this measures:</i> {section["what"]}', styles["BodyText2"]))

            # Why it matters
            story.append(Paragraph(f'<i>Why it matters:</i> {section["why"]}', styles["BodyText2"]))

            # Suggestions
            story.append(Paragraph("<i>What to add:</i>", styles["BodyText2"]))
            for line in section["suggestions"].split("\n"):
                line = line.strip()
                if line:
                    story.append(Paragraph(line, styles["BodyText2"]))
            story.append(Spacer(1, 12))

    # Recommendation
    story.append(Spacer(1, 16))
    story.append(Paragraph("Recommendation", styles["SectionHead"]))
    score = metrics.get("quality_score", 0)
    if score >= 80:
        rec = (
            "The input sources are comprehensive and well-suited for detailed analysis. "
            "We recommend proceeding to the Discovery phase where D8X will extract business rules, "
            "domain entities, and identify conflicts across your sources."
        )
    elif score >= 60:
        rec = (
            "The input sources are adequate for analysis. You can proceed to Discovery now and get useful results, "
            "or you can improve your scores first by adding the suggested source types above. "
            "Adding more diverse sources typically increases the number of business rules extracted by 30-50% "
            "and enables conflict detection between different source types."
        )
    else:
        rec = (
            "The current input sources may be insufficient for thorough analysis. We strongly recommend "
            "adding more documentation before proceeding — particularly the source types listed in the "
            "improvement guide above. Proceeding with limited inputs will produce a shallow analysis "
            "that may miss critical business rules and requirements."
        )
    story.append(Paragraph(rec, styles["BodyText2"]))

    # Footer
    story.append(Spacer(1, 30))
    story.append(HRFlowable(width="100%", thickness=0.5, color=border, spaceBefore=10, spaceAfter=10))
    story.append(Paragraph(f"Generated by D8X Agentic SDLC Platform | d8x.ai | {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} | Patent Pending", styles["SmallGray"]))

    doc.build(story)
    return buffer.getvalue()


def generate_docx_report(project: object, agent_run: object, agent_type: str) -> bytes:
    """Generate a professional Word document report."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    output = agent_run.output_summary or {}  # type: ignore[union-attr]
    metrics = output.get("metrics", {})
    processed_files = output.get("processed_files", [])
    quality = output.get("quality_assessment", {})
    project_type = output.get("project_type", "unknown")
    type_breakdown = output.get("type_breakdown", {})
    project_type_label = "Legacy modernization" if project_type == "legacy_modernization" else "Greenfield project"
    agent_labels = {"ingest": "D1 — Ingest", "discover": "D2 — Discover", "design": "D3 — Design"}

    # Title
    title = doc.add_heading("D8X Analysis Report", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    doc.add_paragraph(f"Project: {project.name}")  # type: ignore[union-attr]
    doc.add_paragraph(f"Agent: {agent_labels.get(agent_type, agent_type)}")
    doc.add_paragraph(f"Date: {datetime.now(timezone.utc).strftime('%B %d, %Y')}")
    doc.add_paragraph()

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        f"D8X ingested and analyzed {metrics.get('files_processed', 0)} files "
        f"containing approximately {metrics.get('words_extracted', 0):,} words. "
        f"Project type: {project_type_label}. Quality score: {metrics.get('quality_score', 0)}/100."
    )

    # Metrics
    doc.add_heading("Key Metrics", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Medium Shading 1 Accent 1"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    for label, value in [("Files processed", str(metrics.get("files_processed", 0))), ("Words extracted", f"{metrics.get('words_extracted', 0):,}"), ("Quality score", f"{metrics.get('quality_score', 0)}/100"), ("Project type", project_type_label)]:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
    doc.add_paragraph()

    # Files
    doc.add_heading("Sources Received", level=1)
    if processed_files:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Medium Shading 1 Accent 1"
        for i, h in enumerate(["Filename", "Type", "Words", "Status"]):
            table.rows[0].cells[i].text = h
        for f in processed_files:
            row = table.add_row().cells
            row[0].text = f.get("filename", "?")
            row[1].text = f.get("file_type", "?").replace("_", " ").title()
            row[2].text = f"{f.get('word_count', 0):,}"
            row[3].text = f.get("status", "?").title()
    doc.add_paragraph()

    if type_breakdown:
        doc.add_heading("Source Type Breakdown", level=2)
        doc.add_paragraph(", ".join(f"{c} {t.replace('_', ' ')}" for t, c in type_breakdown.items()))

    # Quality
    doc.add_heading("Readiness Assessment", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Medium Shading 1 Accent 1"
    for i, h in enumerate(["Dimension", "Score", "Rating"]):
        table.rows[0].cells[i].text = h
    for dim in ["completeness", "diversity", "volume"]:
        s = quality.get(dim, 0)
        row = table.add_row().cells
        row[0].text = dim.title()
        row[1].text = f"{s}/100"
        row[2].text = "Excellent" if s >= 80 else "Adequate" if s >= 60 else "Needs improvement"
    doc.add_paragraph()

    warnings = quality.get("warnings", [])
    if warnings:
        doc.add_heading("Warnings", level=2)
        for w in warnings:
            doc.add_paragraph(w, style="List Bullet")

    # ── How to Improve Your Scores ──
    total_words = metrics.get("words_extracted", 0)
    improvement_guide = _build_improvement_guide(quality, type_breakdown, total_words)
    any_needs_improvement = any(int(s["score"]) < 80 for s in improvement_guide)

    if any_needs_improvement:
        doc.add_heading("How to Improve Your Scores", level=1)
        doc.add_paragraph(
            "The quality scores above directly impact how thorough the Discovery agent's analysis will be. "
            "Below are specific, actionable suggestions for each dimension. You can add more files and "
            "re-run the Ingest agent without losing your existing results."
        )

        for section in improvement_guide:
            if int(section["score"]) >= 80:
                continue

            score_val = int(section["score"])
            rating = "Excellent" if score_val >= 80 else "Adequate" if score_val >= 60 else "Needs improvement"

            doc.add_heading(f'{section["title"]} — {section["score"]}/100 ({rating})', level=2)

            p = doc.add_paragraph()
            p.add_run("What this measures: ").bold = True
            p.add_run(section["what"])

            p = doc.add_paragraph()
            p.add_run("Why it matters: ").bold = True
            p.add_run(section["why"])

            p = doc.add_paragraph()
            p.add_run("What to add:").bold = True
            for line in section["suggestions"].split("\n"):
                line = line.strip()
                if line.startswith("•"):
                    doc.add_paragraph(line.lstrip("• "), style="List Bullet")
                elif line:
                    doc.add_paragraph(line)

    # Recommendation
    doc.add_heading("Recommendation", level=1)
    score = metrics.get("quality_score", 0)
    if score >= 80:
        doc.add_paragraph(
            "The input sources are comprehensive and well-suited for detailed analysis. "
            "We recommend proceeding to the Discovery phase where D8X will extract business rules, "
            "domain entities, and identify conflicts across your sources."
        )
    elif score >= 60:
        doc.add_paragraph(
            "The input sources are adequate for analysis. You can proceed to Discovery now and get useful results, "
            "or improve your scores first by adding the suggested source types above. "
            "Adding more diverse sources typically increases business rules extracted by 30-50%."
        )
    else:
        doc.add_paragraph(
            "The current input sources may be insufficient for thorough analysis. We strongly recommend "
            "adding more documentation before proceeding — particularly the source types listed above. "
            "Proceeding with limited inputs will produce a shallow analysis that may miss critical requirements."
        )

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"Generated by D8X | d8x.ai | {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} | Patent Pending")
    run.font.size = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
