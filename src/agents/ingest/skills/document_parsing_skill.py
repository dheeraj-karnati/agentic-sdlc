"""DocumentParsingSkill: parses PDF, DOCX, XLSX, TXT, MD, HTML files."""

from __future__ import annotations

import logging
from pathlib import Path

from pydantic import BaseModel, Field

from src.agents.base.skill import BaseSkill

logger = logging.getLogger(__name__)


class DocumentSection(BaseModel):
    heading: str = ""
    content: str = ""
    page_number: int = 0
    section_type: str = "text"  # text, heading, list, code_block


class DocumentTable(BaseModel):
    title: str = ""
    headers: list[str] = Field(default_factory=list)
    rows: list[list[str]] = Field(default_factory=list)
    page_number: int = 0


class DocumentMetadata(BaseModel):
    author: str = ""
    created_date: str = ""
    page_count: int = 0
    word_count: int = 0
    file_type: str = ""


class DocumentParsingInput(BaseModel):
    file_path: str
    file_type: str = ""  # pdf, docx, xlsx, txt, md, html — auto-detected if empty


class ParsedDocument(BaseModel):
    sections: list[DocumentSection] = Field(default_factory=list)
    tables: list[DocumentTable] = Field(default_factory=list)
    metadata: DocumentMetadata = Field(default_factory=DocumentMetadata)
    raw_text: str = ""


class DocumentParsingSkill(BaseSkill[DocumentParsingInput, ParsedDocument]):
    """Parses documents preserving structure, headings, tables, and metadata."""

    name = "document_parsing"
    description = "Parse PDF, DOCX, XLSX, TXT, MD, HTML files preserving structure"
    input_model = DocumentParsingInput
    output_model = ParsedDocument

    async def execute(self, input_data: DocumentParsingInput) -> ParsedDocument:
        path = Path(input_data.file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        file_type = input_data.file_type or path.suffix.lstrip(".").lower()

        if file_type == "pdf":
            return self._parse_pdf(path)
        elif file_type == "docx":
            return self._parse_docx(path)
        elif file_type == "xlsx":
            return self._parse_xlsx(path)
        elif file_type in ("txt", "md", "html"):
            return self._parse_text(path, file_type)
        else:
            return self._parse_text(path, file_type)

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        try:
            import pymupdf
        except ImportError:
            logger.warning("pymupdf not installed — reading as raw text")
            return self._parse_text(path, "txt")

        doc = pymupdf.open(str(path))
        sections: list[DocumentSection] = []
        tables: list[DocumentTable] = []
        raw_parts: list[str] = []

        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            raw_parts.append(text)
            for block in page.get_text("dict").get("blocks", []):
                if block.get("type") == 0:  # text block
                    lines = block.get("lines", [])
                    block_text = " ".join(
                        span.get("text", "")
                        for line in lines
                        for span in line.get("spans", [])
                    ).strip()
                    if block_text:
                        # Detect headings by font size
                        max_size = max(
                            (span.get("size", 12) for line in lines for span in line.get("spans", [])),
                            default=12,
                        )
                        sec_type = "heading" if max_size > 14 else "text"
                        sections.append(DocumentSection(
                            heading=block_text[:100] if sec_type == "heading" else "",
                            content=block_text,
                            page_number=page_num,
                            section_type=sec_type,
                        ))

            # Table extraction
            for table in page.find_tables():
                extracted = table.extract()
                if extracted and len(extracted) > 1:
                    tables.append(DocumentTable(
                        headers=[str(c) for c in extracted[0]],
                        rows=[[str(c) for c in row] for row in extracted[1:]],
                        page_number=page_num,
                    ))

        raw_text = "\n\n".join(raw_parts)
        return ParsedDocument(
            sections=sections,
            tables=tables,
            metadata=DocumentMetadata(
                page_count=len(doc),
                word_count=len(raw_text.split()),
                file_type="pdf",
                author=doc.metadata.get("author", "") if doc.metadata else "",
            ),
            raw_text=raw_text,
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed — reading as raw text")
            return self._parse_text(path, "txt")

        doc = Document(str(path))
        sections: list[DocumentSection] = []
        tables: list[DocumentTable] = []
        raw_parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            raw_parts.append(text)
            is_heading = para.style.name.startswith("Heading") if para.style else False
            sections.append(DocumentSection(
                heading=text if is_heading else "",
                content=text,
                section_type="heading" if is_heading else "text",
            ))

        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                tables.append(DocumentTable(
                    headers=rows[0] if rows else [],
                    rows=rows[1:] if len(rows) > 1 else [],
                ))

        raw_text = "\n".join(raw_parts)
        return ParsedDocument(
            sections=sections,
            tables=tables,
            metadata=DocumentMetadata(
                word_count=len(raw_text.split()),
                file_type="docx",
            ),
            raw_text=raw_text,
        )

    def _parse_xlsx(self, path: Path) -> ParsedDocument:
        try:
            from openpyxl import load_workbook
        except ImportError:
            logger.warning("openpyxl not installed")
            return ParsedDocument(metadata=DocumentMetadata(file_type="xlsx"))

        wb = load_workbook(str(path), data_only=True)
        tables: list[DocumentTable] = []
        raw_parts: list[str] = []

        for sheet in wb.sheetnames:
            ws = wb[sheet]
            rows = [[str(cell.value or "") for cell in row] for row in ws.iter_rows()]
            if rows:
                tables.append(DocumentTable(
                    title=sheet,
                    headers=rows[0],
                    rows=rows[1:50],  # limit to 50 data rows
                ))
                raw_parts.append(f"Sheet: {sheet}\n" + "\n".join(
                    "\t".join(row) for row in rows[:20]
                ))

        raw_text = "\n\n".join(raw_parts)
        return ParsedDocument(
            tables=tables,
            metadata=DocumentMetadata(
                word_count=len(raw_text.split()),
                file_type="xlsx",
            ),
            raw_text=raw_text,
        )

    def _parse_text(self, path: Path, file_type: str) -> ParsedDocument:
        text = path.read_text(errors="replace")
        sections: list[DocumentSection] = []
        current_heading = ""

        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            # Detect markdown headings
            if stripped.startswith("#"):
                current_heading = stripped.lstrip("# ").strip()
                sections.append(DocumentSection(
                    heading=current_heading,
                    content=stripped,
                    section_type="heading",
                ))
            else:
                sections.append(DocumentSection(
                    heading=current_heading,
                    content=stripped,
                    section_type="text",
                ))

        return ParsedDocument(
            sections=sections,
            metadata=DocumentMetadata(
                word_count=len(text.split()),
                file_type=file_type,
            ),
            raw_text=text,
        )
